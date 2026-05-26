from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


OUT_DIR = Path("output")
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "기하_김주원_양식맞춤_자외선차단제_보고서.docx"
GRAPH_PATH = OUT_DIR / "sunscreen_form_graph.png"


FONT = "함초롬바탕"


def set_run_font(run, size=10, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold


def set_para_format(p, align=None):
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_format(p, align)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_lines(cell, lines, size=10):
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        set_para_format(p)
        run = p.add_run(line)
        set_run_font(run, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_para(cell, text, size=10, bold=False):
    p = cell.add_paragraph()
    set_para_format(p)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_row_height(row, cm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def e_value(x):
    return -12 * (x - 2.2) ** 2 + 98


def s_value(x):
    return -18 * (x - 1.1) ** 2 + 92


def t_value(x):
    return 0.6 * e_value(x) + 0.4 * s_value(x)


def draw_graph():
    width, height = 950, 520
    ml, mr, mt, mb = 80, 40, 45, 65
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    x_min, x_max = 0, 2.6
    y_min, y_max = 40, 105

    def px(x):
        return ml + (x - x_min) / (x_max - x_min) * (width - ml - mr)

    def py(y):
        return height - mb - (y - y_min) / (y_max - y_min) * (height - mt - mb)

    for y in range(40, 106, 10):
        draw.line((ml, py(y), width - mr, py(y)), fill=(230, 230, 230), width=1)
        draw.text((25, py(y) - 7), str(y), fill=(80, 80, 80))
    for i in range(0, 14):
        x = i * 0.2
        draw.line((px(x), mt, px(x), height - mb), fill=(242, 242, 242), width=1)
        if i % 2 == 0:
            draw.text((px(x) - 11, height - mb + 10), f"{x:.1f}", fill=(80, 80, 80))

    draw.line((ml, mt, ml, height - mb), fill=(40, 40, 40), width=2)
    draw.line((ml, height - mb, width - mr, height - mb), fill=(40, 40, 40), width=2)

    series = [
        ("E(x)", e_value, (35, 99, 175)),
        ("S(x)", s_value, (218, 124, 48)),
        ("T(x)", t_value, (45, 140, 85)),
    ]
    for _, fn, color in series:
        pts = []
        for n in range(261):
            x = n / 100
            y = max(y_min, min(y_max, fn(x)))
            pts.append((px(x), py(y)))
        draw.line(pts, fill=color, width=4)

    ox = 1.65
    oy = t_value(ox)
    draw.line((px(ox), py(oy), px(ox), height - mb), fill=(45, 140, 85), width=2)
    draw.ellipse((px(ox) - 6, py(oy) - 6, px(ox) + 6, py(oy) + 6), fill=(45, 140, 85))
    draw.text((px(ox) + 10, py(oy) - 22), "x=1.65", fill=(45, 100, 65))
    draw.text((ml, 18), "Score model by sunscreen amount", fill=(20, 20, 20))
    draw.text((width // 2 - 75, height - 28), "x (mg/cm^2)", fill=(60, 60, 60))

    lx, ly = width - 230, 65
    for idx, (label, _, color) in enumerate(series):
        y = ly + idx * 25
        draw.line((lx, y, lx + 35, y), fill=color, width=4)
        draw.text((lx + 45, y - 8), label, fill=(60, 60, 60))

    img.save(GRAPH_PATH)


def add_section_table(doc, title, hint, body_lines, min_height_cm):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    row = table.rows[0]
    set_row_height(row, min_height_cm)
    left, right = row.cells
    set_cell_width(left, 3.0)
    set_cell_width(right, 14.1)
    set_cell_shading(left, "F2F2F2")
    add_lines(left, [title, hint] if hint else [title], size=9)
    add_lines(right, body_lines, size=10)
    return table


def build_doc():
    draw_graph()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.line_spacing = 1.6
    styles["Normal"].paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("2026학년도 온양여자고등학교 2학년 기하")
    set_run_font(r, 10)

    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.style = "Table Grid"
    set_row_height(title_table.rows[0], 1.2)
    set_cell_text(title_table.cell(0, 0), "기하 그래프 표현과 탐구 보고서", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("2학년    반    번   이름 : 김주원")
    set_run_font(r, 10)

    add_section_table(
        doc,
        "1. 관련자료 탐색",
        "(도서, 논문, 기사 등의 제목)",
        [
            "식품의약품안전처, 「식약처, 여름 휴가철 식·의약품 건강안전 정보 제공」",
            "식품의약품안전처, 「자외선 차단제 올바른 사용법 리플렛」",
            "대한화장품협회, 「소비자를 위한 화장품 상식: 자외선차단제」",
        ],
        2.4,
    )

    add_section_table(
        doc,
        "2. 관련자료 요약",
        "(10줄이상)",
        [
            "1. 식품의약품안전처 자료는 SPF가 자외선B 차단 정도를 나타내는 지수라고 설명한다.",
            "2. 같은 자료에서 PA는 자외선A 차단 정도를 등급으로 나타내는 표시라고 정리한다.",
            "3. SPF 30 정도에서는 약 95% 이상의 자외선을 차단한다고 소개되어 있다.",
            "4. 따라서 자외선 차단제를 고를 때는 SPF와 PA를 함께 확인해야 한다.",
            "5. 대한화장품협회 자료는 자외선 차단제의 시험 기준 도포량을 2 mg/cm²로 제시한다.",
            "6. 이 기준은 제품의 표시 지수가 일정한 조건에서 측정된 값임을 보여 준다.",
            "7. 실제 사용에서는 표시 지수와 피부에 바르는 양이 함께 차단 효과에 영향을 준다.",
            "8. 도포량이 부족하면 제품의 차단 능력이 충분히 드러나기 어렵다.",
            "9. 반대로 지나치게 많이 바르면 끈적임, 번들거림, 백탁감 때문에 사용감이 낮아진다.",
            "10. 결국 차단 효과와 사용감을 함께 고려한 도포량의 균형점을 생각할 필요가 있다.",
            "11. 이 관계는 증가와 감소가 함께 나타나는 그래프로 표현하기 적합하다.",
            "12. 특히 이차함수의 포물선은 최댓값을 갖기 때문에 최적점을 찾는 데 활용된다.",
        ],
        6.3,
    )

    add_section_table(
        doc,
        "3. 탐구 주제",
        "",
        ["자외선 차단제 도포량에 따른 차단 효과와 사용감의 최적점 분석"],
        1.8,
    )

    add_section_table(
        doc,
        "4. 관련 단원",
        "",
        ["Ⅰ. 이차곡선", "1. 포물선", "이차함수의 그래프, 꼭짓점, 최댓값 해석"],
        2.2,
    )

    add_section_table(
        doc,
        "5. 주제 선정 동기",
        "(5줄이상)",
        [
            "자외선 차단제는 화장품학 진로와 관련이 큰 제품이다.",
            "평소에는 SPF와 PA 같은 표시 지수를 보고 제품의 성능을 판단하기 쉽다.",
            "하지만 실제 피부에 바르는 양이 충분하지 않으면 표시된 차단 효과가 그대로 나타나기 어렵다.",
            "또 도포량을 늘리면 보호 효과는 커질 수 있지만 사용감이 나빠져 다시 적게 바르게 된다.",
            "이처럼 차단 효과와 사용감은 서로 영향을 주는 두 조건이다.",
            "기하에서 배운 포물선의 꼭짓점 개념을 이용하면 두 조건을 동시에 고려한 최적점을 설명하게 된다고 보았다.",
            "그래서 도포량을 변수로 두고 차단 효과와 사용감의 균형을 수학적으로 모델링해 보고자 한다.",
        ],
        5.0,
    )

    doc.add_page_break()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_row_height(table.rows[0], 21.0)
    left, right = table.rows[0].cells
    set_cell_width(left, 3.0)
    set_cell_width(right, 14.1)
    set_cell_shading(left, "F2F2F2")
    add_lines(left, ["6. 탐구 내용", "(단원과 관련지어 설명)", "<참고>", "- 관련 사진 첨부 가능", "(15줄이상)"], size=9)
    right.text = ""
    exploration_lines = [
        "도포량을 x mg/cm²라고 두고, 차단 효과 점수를 E(x), 사용감 점수를 S(x), 종합 점수를 T(x)로 두었다.",
        "대한화장품협회가 제시한 시험 기준 도포량 2 mg/cm²를 기준으로 삼아 차단 효과는 2 mg/cm² 부근에서 높아지도록 설정하였다.",
        "사용감은 많이 바를수록 항상 좋아지는 값이 아니므로, 1 mg/cm² 부근에서 높고 이후 점차 낮아지는 형태로 가정하였다.",
        "차단 효과 점수는 E(x) = -12(x - 2.2)² + 98로 두었다.",
        "이 식은 꼭짓점이 (2.2, 98)인 아래로 열린 포물선이다.",
        "도포량이 2 mg/cm²에 가까워질수록 차단 효과가 높아지는 경향을 나타낸다.",
        "사용감 점수는 S(x) = -18(x - 1.1)² + 92로 두었다.",
        "이 식은 꼭짓점이 (1.1, 92)인 아래로 열린 포물선이다.",
        "도포량이 너무 많아지면 끈적임과 백탁감 때문에 사용감이 떨어지는 상황을 반영한다.",
        "종합 점수는 차단 효과를 더 중요하게 보아 T(x) = 0.6E(x) + 0.4S(x)로 정했다.",
        "식을 정리하면 T(x) = -14.4x² + 47.52x + 52.04가 된다.",
        "이 함수도 아래로 열린 포물선이므로 꼭짓점에서 최댓값을 가진다.",
        "꼭짓점의 x좌표는 x = -b / 2a = -47.52 / (2 × -14.4) = 1.65이다.",
        "따라서 이 모델에서 차단 효과와 사용감을 함께 고려한 최적 도포량은 약 1.65 mg/cm²로 나타난다.",
        "2 mg/cm²는 시험 기준에 가까운 값이지만, 실제 사용자가 끈적임 때문에 충분한 양을 바르지 않는 상황도 고려해야 한다.",
        "1.65 mg/cm²는 차단 효과를 크게 낮추지 않으면서 사용감 감소를 줄이는 균형점으로 해석된다.",
        "아래 표는 여러 도포량에서 E(x), S(x), T(x)를 계산한 값이다.",
    ]
    add_lines(right, exploration_lines, size=10)
    add_para(right, "계산값: x=0.50 → E=63.3, S=85.5, T=72.2", size=10)
    add_para(right, "계산값: x=1.00 → E=80.7, S=91.8, T=85.2", size=10)
    add_para(right, "계산값: x=1.50 → E=92.1, S=89.1, T=90.9", size=10)
    add_para(right, "계산값: x=1.65 → E=94.4, S=86.6, T=91.2", size=10)
    add_para(right, "계산값: x=2.00 → E=97.5, S=77.4, T=89.5", size=10)
    add_para(right, "계산값: x=2.50 → E=96.9, S=56.7, T=80.8", size=10)

    add_para(right, "표를 보면 T(x)는 x=1.65 부근에서 가장 크고, x=2.0 이후에는 사용감 감소 때문에 종합 점수가 낮아진다.", size=10)
    p = right.add_paragraph()
    set_para_format(p, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(GRAPH_PATH), width=Inches(4.9))
    add_para(right, "그래프에서도 E(x), S(x), T(x)가 각각 포물선 형태로 나타나며, 초록색 종합 점수의 꼭짓점이 최적점이다.", size=10)
    add_para(right, "이 탐구는 실제 제품의 실험값을 그대로 계산한 것은 아니며, 공개 자료의 기준과 사용 상황을 바탕으로 기하 개념을 적용한 수학적 모델이다.", size=10)

    doc.add_page_break()
    add_section_table(
        doc,
        "7. 결론 및 느낀 점",
        "(탐구한 내용을 바탕으로 자신이 하고 싶은 주장 또 탐구 결과를 상세히 기록할 것, 차후 과세특에 활용 예정)",
        [
            "이번 탐구를 통해 자외선 차단제의 효과를 판단할 때 제품의 표시 지수만 보는 태도에서 벗어나 실제 도포량을 함께 고려해야 함을 알게 되었다.",
            "수학적으로는 차단 효과와 사용감을 각각 이차함수로 나타내고, 두 함수를 가중 평균한 종합 점수의 꼭짓점을 구했다.",
            "그 결과 이 모델에서는 x=1.65 mg/cm² 부근이 차단 효과와 사용감이 균형을 이루는 지점으로 나타났다.",
            "화장품학 관점에서는 높은 차단 지수를 가진 제품이라도 사용자가 충분히 바르지 않으면 실제 효과가 낮아진다는 점이 중요하다.",
            "따라서 좋은 자외선 차단제는 성분의 차단 능력과 발림성, 백탁감, 유분감까지 함께 설계되어야 한다.",
            "후속 탐구에서는 실제 제품을 정해 도포량별 사용감 설문과 UV 비즈 반응을 비교하고, 측정값을 이용해 더 정확한 포물선 모델을 만들고 싶다.",
            "이 활동을 통해 포물선의 꼭짓점이 교과서 속 계산에 머무르지 않고, 제품 사용 조건의 최적점을 설명하는 데에도 쓰인다는 점을 확인했다.",
            "또 수학적 모델링은 현실을 완전히 대신하는 답은 아니며, 여러 조건 사이의 관계를 이해하게 해 주는 분석 방법이라는 점을 배웠다.",
            "앞으로 화장품 제형을 탐구할 때도 성분의 기능과 사용자의 실제 행동을 함께 고려하는 태도를 갖고 싶다.",
            "특히 안정성, 제형, 사용감처럼 진로와 관련된 문제를 수학적 그래프로 표현하는 탐구를 계속 확장하고 싶다.",
        ],
        8.0,
    )

    add_section_table(
        doc,
        "8. 참고문헌",
        "(반드시 자료에 대한 출처를 밝힐 것, 경로 URL로 표기할 것, 3개이상)",
        [
            "1. 식품의약품안전처, 「식약처, 여름 휴가철 식·의약품 건강안전 정보 제공」, 2017.07.20., https://impfood.mfds.go.kr/CFBBB02F02/getCntntsDetail?cntntsSn=282450",
            "2. 식품의약품안전처, 「자외선 차단제 올바른 사용법 리플렛」, 2014.05.30., https://www.mfds.go.kr/brd/m_641/view.do?seq=17224",
            "3. 대한화장품협회, 「소비자를 위한 화장품 상식: 자외선차단제」, https://kcia.or.kr/pedia/sub03/sub03_01.php?no=303&type=view",
        ],
        4.2,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
