from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_DIR = Path("output")
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "기하_김주원_자외선차단제_최적점_보고서.docx"
GRAPH_PATH = OUT_DIR / "sunscreen_model_graph.png"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text, style=None, bold_head=None):
    p = doc.add_paragraph(style=style)
    if bold_head and text.startswith(bold_head):
        r1 = p.add_run(bold_head)
        r1.bold = True
        r2 = p.add_run(text[len(bold_head):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10.5)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Title", 22, "000000", 0, 8),
        ("Heading 1", 15, "000000", 14, 6),
        ("Heading 2", 12, "434343", 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def e_value(x):
    return -12 * (x - 2.2) ** 2 + 98


def s_value(x):
    return -18 * (x - 1.1) ** 2 + 92


def t_value(x):
    return 0.6 * e_value(x) + 0.4 * s_value(x)


def draw_graph():
    width, height = 1200, 720
    margin_left, margin_right = 105, 50
    margin_top, margin_bottom = 55, 90
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    x_min, x_max = 0, 2.6
    y_min, y_max = 40, 105

    def px(x):
        return margin_left + (x - x_min) / (x_max - x_min) * (width - margin_left - margin_right)

    def py(y):
        return height - margin_bottom - (y - y_min) / (y_max - y_min) * (height - margin_top - margin_bottom)

    for y in range(40, 106, 10):
        draw.line((margin_left, py(y), width - margin_right, py(y)), fill=(230, 230, 230), width=1)
        draw.text((25, py(y) - 8), str(y), fill=(80, 80, 80))
    for i in range(0, 14):
        x = i * 0.2
        draw.line((px(x), margin_top, px(x), height - margin_bottom), fill=(242, 242, 242), width=1)
        if i % 2 == 0:
            draw.text((px(x) - 12, height - margin_bottom + 12), f"{x:.1f}", fill=(80, 80, 80))

    draw.line((margin_left, margin_top, margin_left, height - margin_bottom), fill=(40, 40, 40), width=2)
    draw.line((margin_left, height - margin_bottom, width - margin_right, height - margin_bottom), fill=(40, 40, 40), width=2)

    series = [
        ("Protection E(x)", e_value, (35, 99, 175)),
        ("Feel S(x)", s_value, (218, 124, 48)),
        ("Balance T(x)", t_value, (45, 140, 85)),
    ]
    for label, fn, color in series:
        points = []
        for n in range(0, 261):
            x = n / 100
            y = max(y_min, min(y_max, fn(x)))
            points.append((px(x), py(y)))
        draw.line(points, fill=color, width=4)

    optimum_x = 1.65
    optimum_y = t_value(optimum_x)
    draw.ellipse((px(optimum_x) - 7, py(optimum_y) - 7, px(optimum_x) + 7, py(optimum_y) + 7), fill=(45, 140, 85))
    draw.line((px(optimum_x), py(optimum_y), px(optimum_x), height - margin_bottom), fill=(45, 140, 85), width=2)
    draw.text((px(optimum_x) + 12, py(optimum_y) - 24), "x = 1.65", fill=(45, 100, 65))

    draw.text((margin_left, 20), "Sunscreen amount model", fill=(20, 20, 20))
    draw.text((width // 2 - 95, height - 45), "amount x (mg/cm^2)", fill=(60, 60, 60))
    draw.text((15, 28), "score", fill=(60, 60, 60))

    legend_x, legend_y = width - 300, 75
    for idx, (label, _, color) in enumerate(series):
        y = legend_y + idx * 32
        draw.line((legend_x, y, legend_x + 42, y), fill=color, width=5)
        draw.text((legend_x + 54, y - 10), label, fill=(60, 60, 60))

    img.save(GRAPH_PATH)


def build_doc():
    draw_graph()
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("자외선 차단제 도포량에 따른 차단 효과와 사용감의 최적점 분석")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("기하 | 2학년 | 김주원 | 화장품학 진로 탐구")
    meta_run.font.name = "Arial"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading("1. 탐구 동기", level=1)
    add_paragraph(
        doc,
        "자외선 차단제는 표시된 SPF와 PA 지수만 보고 고르는 경우가 많지만, 실제 피부에 바르는 양과 사용감도 효과를 좌우한다. "
        "식품의약품안전처 자료에 따르면 SPF는 자외선B 차단 정도, PA는 자외선A 차단 정도를 나타내며, SPF 30 정도에서 약 95% 이상의 자외선이 차단된다고 설명한다. "
        "또한 대한화장품협회 자료는 자외선 차단제의 시험 기준 도포량을 신체 표면적 1 cm²당 2 mg으로 제시한다. "
        "이 기준에 가까워질수록 차단 효과는 커지지만, 너무 많이 바르면 끈적임과 백탁감 때문에 실제 사용 지속성이 낮아진다. "
        "따라서 본 탐구에서는 도포량을 변수로 두고 차단 효과와 사용감을 함께 고려하여 어느 지점에서 균형이 생기는지 기하적으로 분석하고자 한다."
    )

    doc.add_heading("2. 이론적 배경", level=1)
    add_paragraph(
        doc,
        "SPF와 PA: SPF는 UVB로 인한 홍반을 기준으로 정해지는 차단 지수이고, PA는 UVA 차단 정도를 등급으로 나타낸다. "
        "두 지표는 모두 제품을 일정한 조건에서 바른 뒤 얻은 값이므로, 실제 사용에서는 도포량과 덧바르기 여부가 함께 고려되어야 한다.",
        bold_head="SPF와 PA:",
    )
    add_paragraph(
        doc,
        "도포량 기준: 대한화장품협회는 자외선 차단제의 권장 도포량을 2 mg/cm²로 설명한다. "
        "이 값은 제품 시험에서 유효한 양이므로, 본 탐구에서는 x축의 기준점으로 사용한다. "
        "다만 학생 보고서의 모델은 실제 임상 시험값을 대신하는 값은 아니며, 도포량 변화에 따른 경향을 기하적으로 설명하기 위한 가정이다.",
        bold_head="도포량 기준:",
    )
    add_paragraph(
        doc,
        "이차함수와 포물선: 어떤 양이 증가하다가 일정 지점 이후 증가폭이 줄거나 만족도가 감소하는 상황은 포물선으로 표현하기 쉽다. "
        "포물선의 꼭짓점은 최댓값 또는 최솟값을 나타내므로, 여러 조건을 합친 점수의 꼭짓점을 찾으면 탐구 주제의 최적점을 수학적으로 설명하게 된다.",
        bold_head="이차함수와 포물선:",
    )

    doc.add_heading("3. 수학적 모델 설정", level=1)
    add_paragraph(
        doc,
        "도포량을 x mg/cm²라고 두었다. 차단 효과 점수 E(x)는 2 mg/cm² 부근에서 높아지는 곡선으로, 사용감 점수 S(x)는 지나치게 많이 바르면 낮아지는 곡선으로 잡았다. "
        "차단 효과를 조금 더 중요하게 보기 위해 종합 점수 T(x)는 차단 효과 60%, 사용감 40%의 가중 평균으로 정했다."
    )
    add_paragraph(doc, "E(x) = -12(x - 2.2)² + 98")
    add_paragraph(doc, "S(x) = -18(x - 1.1)² + 92")
    add_paragraph(doc, "T(x) = 0.6E(x) + 0.4S(x) = -14.4x² + 47.52x + 52.04")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["도포량 x", "차단 효과 E(x)", "사용감 S(x)", "종합 점수 T(x)"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "F1F3F4")
    for x in [0.5, 1.0, 1.5, 1.65, 2.0, 2.5]:
        row = table.add_row().cells
        values = [f"{x:.2f}", f"{e_value(x):.1f}", f"{s_value(x):.1f}", f"{t_value(x):.1f}"]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)

    doc.add_paragraph()
    doc.add_picture(str(GRAPH_PATH), width=Inches(6.2))
    cap = doc.add_paragraph("그림 1. 도포량에 따른 차단 효과, 사용감, 종합 점수의 포물선 모델")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("4. 최적점 계산과 해석", level=1)
    add_paragraph(
        doc,
        "종합 점수 T(x)는 아래로 열린 이차함수이다. 아래로 열린 포물선에서는 꼭짓점이 최댓값을 가지므로, 꼭짓점의 x좌표가 차단 효과와 사용감을 함께 고려한 최적 도포량이 된다."
    )
    add_paragraph(doc, "x = -b / 2a = -47.52 / (2 × -14.4) = 1.65")
    add_paragraph(
        doc,
        "계산 결과 종합 점수는 x = 1.65 mg/cm² 부근에서 가장 크게 나타났다. "
        "이 값은 시험 기준인 2 mg/cm²보다 작지만, 차단 효과를 크게 희생하지 않으면서 사용감 감소를 줄이는 지점으로 해석된다. "
        "즉 실제 권장량은 2 mg/cm²에 가깝게 바르는 것이지만, 사용자가 끈적임 때문에 양을 지나치게 줄이는 상황을 생각하면 1.65 mg/cm² 부근은 현실적인 균형점으로 설명된다. "
        "이 분석은 제품의 표시 지수를 바꾸는 계산으로 보기보다, 도포량과 사용 지속성 사이의 관계를 수학적으로 이해하기 위한 모델이다."
    )

    doc.add_heading("5. 화장품학 관점의 해석", level=1)
    add_paragraph(
        doc,
        "화장품 개발에서는 성분의 차단 능력만큼 제형의 발림성, 백탁감, 유분감, 지속성도 중요하다. "
        "같은 SPF 제품이라도 사용자가 충분한 양을 바르지 않으면 기대한 효과가 줄어들기 때문에, 제형 설계는 수학적 최적화 문제와 닮아 있다. "
        "예를 들어 차단 효과를 높이는 원료 함량을 늘리면 보호 점수는 올라가지만, 사용감이 떨어지면 실제 사용량이 감소한다. "
        "따라서 좋은 자외선 차단제는 높은 차단 지수와 편안한 사용감이 함께 유지되도록 설계되어야 한다."
    )

    doc.add_heading("6. 한계와 후속 탐구", level=1)
    add_paragraph(
        doc,
        "본 탐구의 함수식은 실제 실험 자료를 직접 측정한 값은 아니며 공개 자료의 기준과 생활 속 사용 상황을 바탕으로 만든 수학적 모델이다. "
        "따라서 제품 종류, 피부 타입, 날씨, 땀, 물 접촉, 덧바르기 간격에 따라 실제 결과는 달라진다. "
        "후속 탐구에서는 동일한 SPF 제품을 정해 도포량별 백탁감, 끈적임, 재도포 의향을 설문으로 조사하고, 자외선 측정 카드나 UV 비즈를 활용해 차단 경향을 비교하면 더 설득력 있는 그래프가 완성된다."
    )

    doc.add_heading("7. 참고자료", level=1)
    sources = [
        "식품의약품안전처, 「식약처, 여름 휴가철 식·의약품 건강안전 정보 제공」, 2017.07.20.",
        "식품의약품안전처, 「자외선 차단제 올바른 사용법 리플렛」, 2014.05.30.",
        "대한화장품협회, 「소비자를 위한 화장품 상식: 자외선차단제」.",
    ]
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(source)
        for run in p.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            run.font.size = Pt(10)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
