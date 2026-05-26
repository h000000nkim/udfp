from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD = ROOT / "draft_report_revised.md"
OUT = ROOT / "outputs" / "임운택_자율활동_AFE_실험보고서_수정본.docx"


def font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = b.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            b.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "DADCE0")


def para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        r = p.add_run(part[1:-1] if part.startswith("`") else part)
        font(r)
        if part.startswith("`"):
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            r.font.size = Pt(10)
    return p


def table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    if not rows:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    borders(t)
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = t.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            font(r, size=9.5, bold=(ri == 0))
            if ri == 0:
                shade(cell, "F2F4F7")
    doc.add_paragraph()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    for name, size, before, after, color in [
        ("Title", 24, 0, 8, "000000"),
        ("Heading 1", 18, 18, 6, "000000"),
        ("Heading 2", 15, 14, 5, "000000"),
        ("Heading 3", 13, 12, 4, "434343"),
    ]:
        s = styles[name]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.15

    buf = []
    skip = False
    for raw in MD.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("## 작성자 검토 메모"):
            skip = True
        if skip:
            continue
        if line.startswith("|") and line.endswith("|"):
            buf.append(line)
            continue
        if buf:
            table(doc, buf)
            buf = []
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            font(r, size=24)
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("- "):
            p = para(doc, line[2:], "List Bullet")
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.\s+", line):
            p = para(doc, re.sub(r"^\d+\.\s+", "", line), "List Number")
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                caption, path = m.groups()
                img = ROOT / path
                if not img.exists() or not img.is_file():
                    img = Path("/private/tmp") / Path(path).name
                elif path.startswith("attachments/"):
                    tmp_img = Path("/private/tmp") / Path(path).name
                    if tmp_img.exists():
                        img = tmp_img
                if img.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(img), width=Inches(4.8))
                    cap = doc.add_paragraph(caption)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in cap.runs:
                        font(r, size=9, color="555555")
        else:
            para(doc, line)
    if buf:
        table(doc, buf)
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
