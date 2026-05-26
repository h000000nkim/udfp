from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD = ROOT / "draft_report.md"
OUT = ROOT / "outputs" / "임운택_자율활동_AFE_실험보고서_초안.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_in):
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def apply_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph_with_inline_code(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        apply_run_font(run)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            run.font.size = Pt(10)
    return p


def add_markdown_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    col_count = len(rows[0])
    if col_count == 2:
        widths = [1.7, 4.8]
    elif col_count == 3:
        widths = [1.4, 2.55, 2.55]
    elif col_count == 4:
        widths = [1.25, 1.75, 1.75, 1.75]
    else:
        widths = [6.5 / col_count] * col_count
    set_column_widths(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            apply_run_font(run, size=9.5, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
        if r_idx == 0:
            set_repeat_table_header(row=table.rows[0])
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)

    for name, size, before, after, color in [
        ("Title", 24, 0, 8, "000000"),
        ("Heading 1", 18, 18, 6, "000000"),
        ("Heading 2", 15, 14, 5, "000000"),
        ("Heading 3", 13, 12, 4, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    lines = MD.read_text(encoding="utf-8").splitlines()
    table_buf = []
    skip_review = False

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_markdown_table(doc, table_buf)
            table_buf = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("## 작성자 검토 메모"):
            skip_review = True
        if skip_review:
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buf.append(line)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            apply_run_font(run, size=24)
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("- "):
            p = add_paragraph_with_inline_code(doc, line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.\s+", line):
            p = add_paragraph_with_inline_code(doc, re.sub(r"^\d+\.\s+", "", line), style="List Number")
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption, rel_path = match.groups()
                img_path = ROOT / rel_path
                if img_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(4.8))
                    cap = doc.add_paragraph(caption)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        apply_run_font(run, size=9, color="555555")
        else:
            add_paragraph_with_inline_code(doc, line)

    flush_table()
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
