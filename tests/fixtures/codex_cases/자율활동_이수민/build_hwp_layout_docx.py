from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "outputs/자율활동_이수민/final_application_hwp_layout.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=85, start=120, bottom=85, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "555555")


def clear(cell):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    margins(cell)


def run_style(run, size=9, bold=False, color="111111"):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def para(cell, text="", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.08
    if text:
        r = p.add_run(text)
        run_style(r, size=size, bold=bold)
    return p


def label(cell, text, size=9):
    clear(cell)
    shade(cell, "EDEDED")
    p = para(cell, text, bold=True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(0)


def body(cell, texts):
    clear(cell)
    shade(cell, "FFFFFF")
    for item in texts:
        if isinstance(item, tuple):
            para(cell, item[0], bold=item[1])
        else:
            para(cell, item)


def set_grid(row):
    width(row.cells[0], 3.0)
    width(row.cells[1], 3.8)
    width(row.cells[2], 10.3)


def add_wide_row(table, left, texts):
    row = table.add_row()
    set_grid(row)
    label(row.cells[0], left)
    merged = row.cells[1].merge(row.cells[2])
    body(merged, texts)
    return row


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.35)
section.bottom_margin = Cm(1.35)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
style.font.size = Pt(9)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(7)
r = title.add_run("2026학년도 1학년 ‘지역사회 연계 문제해결 탐구 프로젝트’ 신청서")
run_style(r, size=14, bold=True, color="000000")

table = doc.add_table(rows=0, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
borders(table)

add_wide_row(
    table,
    "신청 단위",
    ["모둠", "학번/이름: 제출 전 모둠원 학번과 이름을 학번 순서대로 입력"],
)

add_wide_row(
    table,
    "공통 탐구\n배경",
    [
        ("③ 지자체의 도시침수 대응사업 진행", True),
        "최근 아산시를 포함한 여러 지역에서 집중호우로 인한 도로 침수와 하수 역류 피해가 반복되고 있다. 지자체에서는 우수관로 정비, 빗물받이 확충, 침수 대응 사업을 추진하고 있지만, 아스팔트와 콘크리트가 많은 도시에서는 빗물이 땅속으로 스며들기 어렵다. 그래서 배수 시설을 늘리는 방법만으로는 폭우 때 한꺼번에 몰리는 물을 모두 처리하기 어렵다고 보았다. 우리 모둠은 도시의 물순환이 깨질 때 오염된 물과 수인성 병원체가 어떻게 이동하는지 살펴보고, 빗물 정원처럼 자연의 정화 작용을 활용한 방법이 도움이 되는지 탐구하고자 한다.",
    ],
)

add_wide_row(
    table,
    "탐구 주제\n(세부)",
    [
        "폭우로 인해 깨진 도시의 물순환 과정이 수인성 병원체의 이동과 증식을 어떻게 유발하는지 분석하고, 빗물 정원의 물리적·화학적·생물학적 정화 기능으로 이를 완화하는 방안을 탐구한다.",
    ],
)

add_wide_row(
    table,
    "교과 연계\n작성",
    [
        "이 탐구는 통합과학에서 다룬 물의 순환과 지구 시스템의 상호작용 개념을 활용한다. 폭우는 기권에서 일어나는 강수 현상이며, 많은 비가 짧은 시간에 내리면 수권의 하천과 하수 흐름이 급격히 늘어난다. 도시의 불투수층은 지권으로 스며드는 물의 양을 줄이고, 그 결과 빗물이 지표면을 따라 빠르게 흘러 오염물질을 함께 이동시킨다. 이 과정에서 생물권에 속하는 세균과 수인성 병원체도 물을 따라 확산된다.",
        "빗물 정원은 토양의 여과 작용, 식물의 흡수 작용, 토양 속 미생물의 분해 작용을 이용해 빗물을 천천히 머금고 정화하는 시설이다. 따라서 이 주제는 기권·수권·지권·생물권이 서로 영향을 주고받는 과정을 실제 지역 문제에 적용해 보는 탐구가 된다.",
    ],
)

analysis_rows = []
for sub, texts in [
    (
        "선택 이유",
        [
            "집중호우가 내릴 때 도로가 잠기거나 하수가 역류하면 오염된 물이 주거지와 학교 주변으로 퍼진다. 특히 도시 지역은 포장된 면적이 넓어 빗물이 흙으로 스며들지 못하고 도로 위를 빠르게 흐른다. 이때 빗물은 먼지, 쓰레기, 동물 배설물, 하수의 오염 성분을 함께 운반하고, 수인성 병원체가 넓은 범위로 이동할 가능성도 커진다.",
            "우리 모둠은 침수 문제를 배수 시설 부족만의 문제로 보지 않고, 도시의 물순환이 약해진 결과로 보고자 했다. 또 학교 주변이나 아산시의 도시침수 대응사업과 비교하면서, 빗물 정원 같은 자연 기반 시설이 실제로 어떤 역할을 하는지 알아보고 싶었다.",
        ],
    ),
    (
        "관련된 교과 지식",
        [
            "통합과학의 물질 순환과 지구 시스템 관점에서 보면, 폭우는 여러 권역의 변화를 동시에 일으킨다. 기권에서 많은 비가 내리면 수권의 물의 양이 갑자기 증가하고, 지권에서는 불투수층 때문에 침투량이 줄어든다. 그 결과 지표면 유출량이 커져 오염물질과 병원체가 빠르게 이동한다.",
            "빗물 정원의 정화 과정은 세 가지로 정리된다. 첫째, 자갈과 토양 입자가 물속의 큰 입자를 걸러 내는 물리적 여과가 일어난다. 둘째, 토양의 성분이 일부 오염물질을 붙잡거나 농도를 낮추는 화학적 작용이 일어난다. 셋째, 식물 뿌리 주변과 토양 속 미생물이 유기물을 분해하면서 물의 오염 정도를 낮추는 생물학적 작용이 일어난다. 이 과정을 통해 빗물 정원은 빗물을 천천히 흡수하고 정화하여 도시의 물순환 회복에 도움을 준다.",
        ],
    ),
    (
        "어떤 결과물을\n원하나요?",
        [
            "카드뉴스와 제안 보고서를 함께 제작하고 싶다. 카드뉴스에는 도시침수가 생기는 원인, 물순환 붕괴와 병원체 이동 과정, 빗물 정원의 정화 원리를 그림과 짧은 설명으로 정리한다. 제안 보고서에는 아산시 도시침수 대응사업과 학교 주변 환경을 바탕으로 빗물 정원 설치 가능 장소, 기대 효과, 한계, 후속 탐구 방법을 담는다.",
        ],
    ),
]:
    row = table.add_row()
    set_grid(row)
    analysis_rows.append(row)
    label(row.cells[1], sub)
    body(row.cells[2], texts)

merged_analysis = analysis_rows[0].cells[0].merge(analysis_rows[-1].cells[0])
label(merged_analysis, "선택 주제\n분석", size=9.5)

add_wide_row(
    table,
    "자료\n수집\n계획",
    [
        "지역 통계 자료 및 지자체 공공 데이터 분석: 아산시 강수량, 침수 피해 기사, 도시침수 대응사업 자료를 찾아본다.",
        "학교 주변 지형도 및 로드뷰 분석: 학교 주변의 포장 면적, 빗물 흐름이 모일 만한 지점, 빗물받이 위치를 확인한다.",
        "현장 답사 및 사진 촬영: 학교 주변 배수로 위치와 빗물받이 위치 그리고 물이 고이기 쉬운 장소를 관찰하고 사진으로 기록한다.",
        "관련자 인터뷰: 가능하다면 과학 선생님이나 학교 시설 담당자에게 폭우 때 학교 주변 배수 상황을 질문한다.",
        "기타: 빗물 정원 사례를 찾아 정화 원리와 설치 조건을 정리한다.",
    ],
)

row = table.add_row()
set_grid(row)
label(row.cells[0], "")
merged = row.cells[1].merge(row.cells[2])
label(merged, "활동 계획")

add_wide_row(
    table,
    "최종 결과물\n형태",
    [
        ("카드뉴스 / 제안 보고서", True),
        "카드뉴스는 5~6장 분량으로 구성한다. 1장은 문제 제기, 2장은 도시 물순환이 깨지는 과정, 3장은 수인성 병원체 이동 가능성, 4장은 빗물 정원의 정화 원리, 5장은 학교 주변 또는 아산시에 적용 가능한 제안, 6장은 기대 효과와 후속 탐구로 구성한다.",
        "제안 보고서는 신청서 내용을 바탕으로 탐구 배경, 자료 분석, 교과 개념 적용, 빗물 정원 제안, 한계와 보완점을 문단형으로 정리한다.",
    ],
)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date.paragraph_format.space_before = Pt(8)
r = date.add_run("2026.    .    .")
run_style(r, size=10)

sign = doc.add_table(rows=2, cols=1)
sign.alignment = WD_TABLE_ALIGNMENT.RIGHT
sign.autofit = False
borders(sign)
for row in sign.rows:
    clear(row.cells[0])
    width(row.cells[0], 3.2)
label(sign.rows[0].cells[0], "담당교사 확인")
shade(sign.rows[1].cells[0], "FFFFFF")

doc.save(OUT)
print(OUT)
