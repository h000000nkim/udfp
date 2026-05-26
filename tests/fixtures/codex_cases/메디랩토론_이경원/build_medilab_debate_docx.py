from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


OUT = "이경원_동아리활동_메디랩_토론보고서_양식맞춤.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=140, bottom=80, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_fixed_table_geometry(table, label_twips, body_twips):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(label_twips + body_twips))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in (label_twips, body_twips):
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)


def set_repeatable_font(run, size=10, bold=False):
    run.font.name = "함초롬바탕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "함초롬바탕")
    run.font.size = Pt(size)
    run.bold = bold


def put_text(cell, text, bold=False, align=None, size=10):
    cell.text = ""
    for i, part in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        if align is not None:
            paragraph.alignment = align
        run = paragraph.add_run(part)
        set_repeatable_font(run, size=size, bold=bold)


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(6)
section.bottom_margin = Mm(6)
section.left_margin = Mm(10)
section.right_margin = Mm(10)

styles = doc.styles
styles["Normal"].font.name = "함초롬바탕"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "함초롬바탕")
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
title.paragraph_format.space_before = Pt(0)
title_run = title.add_run("메디랩 토론 보고서")
set_repeatable_font(title_run, size=13, bold=True)

rows = [
    ("학번", "DJ9370"),
    ("이름", "이경원"),
    ("토론 주제", "첨단 반도체 기술에 대한 국가의 수출 통제는 정당하다."),
    (
        "핵심 쟁점",
        "반도체 기술은 안보와 관련되지만, 설계·장비·소재·제조·시장이 국제적으로 나뉘어 움직인다. 쟁점은 수출 통제가 기술 보호보다 공급망 불안과 기업 부담을 더 키우는가이다.",
    ),
    (
        "주장 1",
        "첨단 반도체 기술에 대한 넓은 수출 통제는 정당하지 않다. 반도체는 한 국가가 홀로 완성하는 기술이 아니라 여러 나라의 연구, 장비, 소재, 생산 능력이 함께 작동해야 발전하기 때문이다.",
    ),
    (
        "근거",
        "『반도체 오디세이』는 반도체를 설계·장비·제조·시장 전략이 함께 움직이는 산업으로 설명한다. 전자공학적으로도 회로 설계, 공정 제어, 소재 특성, 열 관리가 함께 개선되어야 성능이 높아진다. 따라서 특정 기술의 이동을 정치적으로 넓게 막으면 관련 연구와 생산이 함께 흔들린다.",
    ),
    (
        "주장 2",
        "수출 통제 강화는 기술 질서를 안정시키기보다 미·중 갈등과 공급망 불안을 키운다. 한국처럼 세계 반도체 공급망에 깊이 참여하는 국가는 기업 활동과 기술 개발에서 더 큰 부담을 받는다.",
    ),
    (
        "근거",
        "대외경제정책연구원은 미·중 반도체 패권 경쟁이 한국 반도체 산업의 공급망 위험에도 영향을 준다고 분석한다. 통제를 받은 국가는 자체 기술 개발에 더 많은 자원을 투입해 장기적으로 새 경쟁자가 될 수 있다. 한국도 장비, 소재, 시스템 반도체 분야에서는 여러 나라의 기술과 시장에 기대므로 수출 통제 확대는 기업 부담으로 이어진다.",
    ),
    (
        "반론",
        "찬성 측은 첨단 반도체가 군사 기술이나 전략 산업에 쓰이므로 국가가 강하게 통제해야 한다고 볼 수 있다. 그러나 안보 목적이 있다고 해서 모든 첨단 기술 이동을 넓게 막는 방식이 정당해지는 것은 아니다. 군사적 전용 가능성이 뚜렷한 범위는 관리하되, 연구 협력과 정상적인 기업 활동까지 위축시키는 통제는 줄여야 한다.",
    ),
    (
        "결론 및 느낀점",
        "나는 첨단 반도체 기술 보호의 필요성은 인정하지만, 국가의 넓은 수출 통제를 정당하다고 보기는 어렵다고 판단했다. 『반도체 오디세이』를 읽으며 반도체가 한 나라의 기술이 아니라 세계적 분업 속에서 발전한 산업임을 알게 되었다. 앞으로 전자공학을 공부할 때 기술 자체뿐 아니라 공급망, 국가 정책, 기업 전략까지 함께 보아야겠다고 느꼈다.",
    ),
]

table = doc.add_table(rows=len(rows), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_borders(table)

label_twips = int(Mm(29.11).twips)
body_twips = int(Mm(158.89).twips)
set_fixed_table_geometry(table, label_twips, body_twips)
heights_mm = [7.2, 7.2, 15.5, 27.0, 27.0, 27.0, 27.0, 27.0, 27.0, 27.0]

for idx, (label, body) in enumerate(rows):
    row = table.rows[idx]
    row.height = Mm(heights_mm[idx])
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    label_cell = row.cells[0]
    body_cell = row.cells[1]
    for cell, width in ((label_cell, label_twips), (body_cell, body_twips)):
        set_cell_width(cell, width)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(label_cell, "EDEDED")
    put_text(label_cell, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    put_text(body_cell, body, size=7.8 if idx in (3, 5, 7) else 8.5)

doc.save(OUT)
print(OUT)
