from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = "outputs/미적분1_김도현"
REPORT_OUT = f"{BASE}/김도현_미적분1_보고서_양식준수.docx"
SELF_OUT = f"{BASE}/김도현_미적분1_자기평가서_양식준수.docx"


REPORT_TITLE = "RC 회로의 충전과 방전 과정에서 나타나는 전압 변화와 미분의 관계 탐구"

INTRO = [
    "전기전자공학에서는 전압과 전류가 시간에 따라 어떻게 변하는지를 이해하는 일이 중요하다고 생각한다. 특히 커패시터가 들어간 회로는 처음에는 빠르게 변하다가 시간이 지날수록 변화가 느려지는 모습을 보인다. 이 현상은 전류가 흐른다는 설명만으로는 부족하고, 시간에 따른 변화율을 함께 보아야 더 정확하게 이해된다.",
    "미적분Ⅰ에서 배운 미분은 어떤 양이 순간적으로 얼마나 빠르게 변하는지를 나타낸다. 그래서 RC 회로의 전압 변화를 미분의 관점으로 해석하면, 전자회로에서 미적분이 맡는 역할을 확인하게 된다고 생각했다. 이번 탐구에서는 저항과 커패시터로 이루어진 RC 회로를 중심으로, 충전과 방전 과정의 전압 변화가 왜 곡선으로 나타나는지, 그리고 그 변화율이 회로의 성질과 어떤 관련을 가지는지를 살펴보고자 한다.",
]

BODY = [
    "미분은 함수의 평균 변화율을 한 점에서의 순간 변화율로 확장한 개념이다. 함수 y=f(x)에서 x가 조금 변할 때 y가 얼마나 변하는지를 살피면 평균 변화율을 구하게 되고, 변화 구간을 점점 작게 하면 한 순간의 변화율인 미분계수에 가까워진다. 도함수는 각 x값에서의 순간 변화율을 함수로 나타낸 것이다.",
    "시간에 따라 변하는 전압을 V(t)라고 하면, V'(t)는 t초에서 전압이 얼마나 빠르게 증가하거나 감소하는지를 뜻한다. 전압이 빠르게 변할수록 도함수의 절댓값은 커지고, 전압 변화가 거의 멈추면 도함수의 값은 0에 가까워진다. 이 해석은 회로의 충전과 방전 현상을 이해하는 데 직접 사용된다.",
    "RC 회로에서 자주 나타나는 함수는 지수함수 꼴의 변화이다. 예를 들어 충전 과정의 전압은 이상적인 상황에서 다음 식으로 나타난다.",
    "V(t)=V0(1-e^(-t/RC))",
    "여기서 V0는 최종적으로 가까워지는 전압, R은 저항, C는 커패시터의 전기용량이다. RC는 시정수라고 하며, 회로가 얼마나 빠르게 안정 상태에 가까워지는지를 나타낸다. 이 식을 미분하면 전압의 증가 속도가 시간이 지날수록 작아진다는 사실이 드러난다. 즉 처음에는 변화율이 크지만, 시간이 지날수록 전압이 최종값에 가까워지면서 변화율이 감소한다.",
    "RC 회로는 저항 R과 커패시터 C로 이루어진 회로이다. 커패시터는 전하를 저장하는 부품이고, 저항은 전류의 흐름을 조절한다. 커패시터가 처음 비어 있을 때 전원을 걸어 주면 전하가 빠르게 저장되기 시작한다. 이때 전압은 처음에는 급격히 증가하지만, 시간이 지날수록 증가 속도가 느려진다.",
    "충전 과정에서 전압을 V(t)=V0(1-e^(-t/RC))라고 두면, 이 함수의 도함수는 다음과 같은 의미를 가진다.",
    "V'(t)=(V0/RC)e^(-t/RC)",
    "이 식에서 e^(-t/RC)는 시간이 지날수록 작아진다. 따라서 V'(t)도 시간이 지날수록 작아진다. 이것은 커패시터가 비어 있을 때는 전하가 빠르게 저장되지만, 점점 채워질수록 전하의 이동 속도가 줄어드는 실제 회로 현상과 맞아떨어진다.",
    "방전 과정도 같은 방식으로 설명된다. 충전된 커패시터가 저항을 통해 방전될 때 전압은 다음과 같이 감소한다.",
    "V(t)=V0e^(-t/RC)",
    "이 함수의 도함수는 다음과 같다.",
    "V'(t)=-(V0/RC)e^(-t/RC)",
    "부호가 음수인 것은 전압이 감소하고 있음을 뜻한다. 또한 시간이 지날수록 e^(-t/RC)가 작아지기 때문에 감소 속도도 점점 느려진다. 즉 방전 초기에는 전압이 빠르게 떨어지지만, 시간이 지나면 전압 변화가 완만해진다.",
    "여기서 중요한 값은 시정수 RC이다. R이나 C가 커지면 RC 값도 커진다. 그러면 지수함수의 감소가 느려져 전압 변화가 더 천천히 일어난다. 반대로 R이나 C가 작으면 RC 값이 작아져 전압이 더 빠르게 변한다. 이처럼 하나의 수식 안에서 회로 부품의 값과 전압 변화의 속도가 함께 설명된다.",
    "예를 들어 같은 전압을 걸어 주더라도 커패시터의 용량이 큰 회로는 전하를 더 많이 저장하므로 충전되는 데 시간이 더 걸린다. 이를 그래프로 나타내면 C가 클수록 전압 곡선이 완만해진다. 미분의 관점에서는 같은 시간 t에서 도함수의 값이 달라지며, 이것이 회로의 반응 속도 차이로 해석된다.",
    "RC 회로는 전기전자 분야에서 다양하게 사용된다. 예를 들어 신호가 갑자기 변할 때 그 변화를 완만하게 만들거나, 특정한 빠르기의 변화만 통과시키는 회로를 설계할 때 RC 회로의 성질이 활용된다. 센서 회로에서도 측정값이 순간적으로 흔들릴 때 이를 부드럽게 처리하는 데 비슷한 원리가 쓰인다.",
    "전자공학에서는 회로가 입력 신호에 얼마나 빠르게 반응하는지가 중요하다. 너무 느리게 반응하면 필요한 정보를 제때 처리하지 못하고, 너무 민감하게 반응하면 불필요한 잡음까지 크게 받아들인다. 따라서 저항과 커패시터의 값을 조절하여 원하는 반응 속도를 만드는 과정이 필요하다. 이때 시정수 RC는 회로 설계자가 참고하는 중요한 기준이 된다.",
    "또한 신호 처리에서는 시간에 따라 변하는 전압을 함수로 보고, 그 변화율을 해석하는 일이 자주 등장한다. 미분은 신호의 급격한 변화나 경향을 파악하는 데 도움이 된다. 이번 탐구에서 다룬 RC 회로는 전자공학의 기본 예시이지만, 시간에 따른 변화와 회로의 성능을 함께 해석한다는 점에서 진로와 관련된 의미가 크다.",
]

CONCLUSION = [
    "이 탐구를 통해 미분이 실제 회로에서 시간에 따라 변하는 물리량의 변화를 설명하는 도구라는 점을 알게 되었다. 특히 RC 회로는 고등학교 수준의 미분 개념으로도 변화율의 의미를 비교적 분명하게 보여 주는 사례이다.",
    "이번 탐구를 하면서 미적분Ⅰ에서 배운 미분 개념이 전기전자공학의 기초 현상을 설명하는 데 쓰인다는 점을 알게 되었다. 처음에는 RC 회로의 식이 외워야 하는 공식처럼 보였지만, 도함수를 통해 살펴보니 전압 변화가 왜 처음에는 빠르고 나중에는 느려지는지 더 자연스럽게 이해되었다.",
    "특히 시정수 RC가 회로의 반응 속도를 결정하는 기준이라는 점이 인상적이었다. 저항이나 커패시터의 값이 바뀌면 그래프의 모양과 변화율이 함께 달라진다. 이 과정에서 수학적 함수와 실제 회로 부품 사이의 관계를 확인했다.",
    "다만 이번 탐구는 RC 회로의 기본적인 충전과 방전 현상을 중심으로 진행했기 때문에, 실제 회로에서 나타나는 오차나 복잡한 신호 변화까지는 충분히 다루지 못했다. 후속 탐구에서는 오실로스코프를 이용해 실제 RC 회로의 전압 변화를 측정하고, 이론식과 실제 측정값을 비교해 보고 싶다. 또한 미분회로와 적분회로에서 입력 신호의 형태가 출력 신호로 어떻게 바뀌는지도 더 알아보고 싶다.",
]

REFERENCES = [
    "과제 첨부 자료, 미적분Ⅰ의 원리와 응용에 관한 심층 탐구 보고서 양식",
    "과제 첨부 자료, 미적분Ⅰ의 원리와 응용에 관한 심층 탐구 자기평가서",
    "고등학교 미적분Ⅰ 교과 개념: 평균 변화율, 순간 변화율, 도함수",
]


SELF_ANSWERS = [
    ("1. 탐구 주제는?", REPORT_TITLE),
    ("2. 왜 이 주제를 탐구했나요? (3줄 내외)", "전기전자공학에서 시간에 따라 변하는 전압과 전류를 이해하는 일이 중요하다고 생각했다. RC 회로의 전압 변화가 미적분Ⅰ에서 배운 순간 변화율과 직접 관련된다는 점이 흥미로워 이 주제를 탐구했다."),
    ("3. 탐구 내용을 핵심 요약해본다면? (4줄 내외)", "RC 회로의 충전 전압은 V(t)=V0(1-e^(-t/RC)), 방전 전압은 V(t)=V0e^(-t/RC)로 나타난다. 이를 미분하면 전압 변화율이 시간이 지날수록 작아짐을 확인했다. 시정수 RC가 클수록 회로의 반응 속도가 느려진다는 점도 알게 되었다."),
    ("4. 어려웠던 점은 무엇이며 어떻게 극복하였나? (5줄 내외)", "처음에는 지수함수 형태의 식이 실제 회로 현상과 어떻게 관련되는지 이해하기 어려웠다. 그래서 전압 그래프의 모양과 도함수의 부호 및 크기를 함께 비교했다. 충전과 방전의 변화율을 각각 해석하면서 수식과 회로 현상을 대응시킬 수 있었다."),
    ("5. 새롭게 배우고 느낀 점은? (2줄 내외)", "미분이 그래프 계산에 머무르지 않고 전자회로의 반응 속도를 설명하는 도구가 된다는 점을 배웠다. 전기전자공학 진로와 미적분Ⅰ 개념이 생각보다 가까이 이어져 있다는 점이 인상적이었다."),
]

SELF_SYNTHESIS = "이번 탐구에서는 RC 회로의 충전과 방전 과정에서 나타나는 전압 변화를 미분의 관점으로 분석했다. 충전 과정에서는 전압이 처음에 빠르게 증가하다가 점점 최종 전압에 가까워지고, 방전 과정에서는 전압이 빠르게 감소하다가 점차 완만해진다. 이를 각각 V(t)=V0(1-e^(-t/RC)), V(t)=V0e^(-t/RC)로 나타내고 도함수를 살펴보니, 전압 변화율이 시간이 지날수록 작아진다는 사실을 확인했다. 특히 시정수 RC가 회로의 반응 속도를 결정한다는 점을 알게 되었다. 저항이나 커패시터의 값이 달라지면 그래프의 모양과 변화율도 함께 달라진다. 이 과정에서 미분이 실제 전기전자 회로의 변화를 설명하는 데 쓰인다는 점을 배웠다. 후속 탐구에서는 실제 RC 회로를 구성해 전압 변화를 측정하고 이론식과 비교해 보고 싶다."


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Gulim"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "굴림")
    style.font.size = Pt(12)
    return doc


def set_cell(cell, text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = "Gulim"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "굴림")
            r.font.size = Pt(size)
            r.font.bold = bold


def borders(cell, color="000000", size="6"):
    pr = cell._tc.get_or_add_tcPr()
    tc_borders = pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = tc_borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def margins(cell, top=80, start=120, bottom=80, end=120):
    pr = cell._tc.get_or_add_tcPr()
    mar = pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        pr.append(mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = mar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table_width(table, pct="5000"):
    pr = table._tbl.tblPr
    tbl_w = pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), pct)


def fmt_p(p, size=12, bold=False, align=None, before=0, after=6, line=1.45):
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for r in p.runs:
        r.font.name = "Gulim"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "굴림")
        r.font.size = Pt(size)
        r.font.bold = bold


def add_para(doc, text="", size=12, bold=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.add_run(text)
    fmt_p(p, size=size, bold=bold, align=align, before=before, after=after)
    return p


def add_report_top(doc):
    table = doc.add_table(rows=3, cols=4)
    table_width(table)
    row = table.rows[0]
    row.cells[0].merge(row.cells[3])
    set_cell(row.cells[0], "<미적분Ⅰ의 원리와 응용에 관한 심층 탐구> 보고서", True, 13, WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ["학번", "21107", "이름", "김도현"],
        ["제목", REPORT_TITLE, "", ""],
    ]
    for r, values in enumerate(rows, start=1):
        if r == 2:
            table.rows[r].cells[1].merge(table.rows[r].cells[3])
        for c, value in enumerate(values):
            if r == 2 and c > 1:
                continue
            set_cell(table.rows[r].cells[c], value, bold=(c in (0, 2) or r == 0), size=11, align=WD_ALIGN_PARAGRAPH.CENTER if c in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT)
    for row in table.rows:
        for cell in row.cells:
            borders(cell)
            margins(cell)


def build_report():
    doc = setup_doc()
    add_report_top(doc)
    add_para(doc, "#글자크기12포인트 #글꼴 굴림 #그림표삽입가능 #2페이지초과작성(채점기준제외)", size=10, after=10)
    sections = [
        ("Ⅰ. 서론 (이 주제를 선택한 이유와 계기, 탐구 내용의 핵심 등을 작성)", INTRO),
        ("Ⅱ. 본론 (탐구 내용을 보고서 형식으로 체계를 갖추어 작성, 그림 삽입 가능함)", BODY),
        ("Ⅲ. 결론 (탐구의 결론과 핵심, 새롭게 배운점, 추가로 탐구하고 싶은 내용 등을 작성)", CONCLUSION),
    ]
    for title, paragraphs in sections:
        add_para(doc, title, bold=True, before=6, after=6)
        for text in paragraphs:
            align = WD_ALIGN_PARAGRAPH.CENTER if text.startswith("V(") else WD_ALIGN_PARAGRAPH.JUSTIFY
            add_para(doc, text, align=align, after=5)

    add_para(doc, "Ⅳ. 참고문헌 (읽은 책, 논문, 참고 누리집 주소 등을 작성)", bold=True, before=6, after=4)
    table = doc.add_table(rows=5, cols=2)
    table_width(table)
    for i in range(5):
        set_cell(table.rows[i].cells[0], f"[{i + 1}]", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.rows[i].cells[1], REFERENCES[i] if i < len(REFERENCES) else "", size=10)
        for cell in table.rows[i].cells:
            borders(cell)
            margins(cell)
    add_para(doc, "※ 참고문헌 입력란이 모자라면 직접 추가해서 작성 가능함", size=10, after=8)
    add_para(doc, "※ 채점 기준", bold=True, before=6, after=4)
    rubric = doc.add_table(rows=5, cols=3)
    table_width(rubric)
    data = [
        ["구분", "기준", "점수"],
        ["탐구활동 과정\n(60점)", "탐구 과정의 논리적 전개가 매우 우수한 경우", "60"],
        ["", "탐구 과정의 논리적 전개가 우수한 경우", "50"],
        ["", "탐구 과정의 논리적 전개가 보통인 경우", "40"],
        ["", "탐구 과정의 논리적 전개가 미흡한 경우", "30"],
    ]
    for r, row in enumerate(rubric.rows):
        for c, cell in enumerate(row.cells):
            set_cell(cell, data[r][c], bold=(r == 0), size=9, align=WD_ALIGN_PARAGRAPH.CENTER if c != 1 else WD_ALIGN_PARAGRAPH.LEFT)
            borders(cell)
            margins(cell)
    doc.save(REPORT_OUT)


def add_self_top(doc):
    table = doc.add_table(rows=2, cols=5)
    table_width(table)
    table.rows[0].cells[0].merge(table.rows[0].cells[4])
    set_cell(table.rows[0].cells[0], "<미적분Ⅰ의 원리와 응용에 관한 심층 탐구> 자기 평가서", True, 13, WD_ALIGN_PARAGRAPH.CENTER)
    values = ["학번", "21107", "이름", "김도현"]
    table.rows[1].cells[0].merge(table.rows[1].cells[1])
    set_cell(table.rows[1].cells[0], values[0], True, 10, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.rows[1].cells[2], values[1], False, 10, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.rows[1].cells[3], values[2], True, 10, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table.rows[1].cells[4], values[3], False, 10, WD_ALIGN_PARAGRAPH.CENTER)
    for row in table.rows:
        for cell in row.cells:
            borders(cell)
            margins(cell)


def build_self_eval():
    doc = setup_doc()
    sec = doc.sections[0]
    sec.left_margin = Cm(1.3)
    sec.right_margin = Cm(1.3)
    add_self_top(doc)

    rubric = doc.add_table(rows=5, cols=3)
    table_width(rubric)
    rows = [
        ["구분", "기준", "점수"],
        ["자기\n평가", "탐구 결과에 대한 자기 환류가 매우 우수한 경우", "20"],
        ["", "탐구 결과에 대한 자기 환류가 우수한 경우", "15"],
        ["", "탐구 결과에 대한 자기 환류가 보통인 경우", "10"],
        ["", "탐구 결과에 대한 자기 환류가 미흡한 경우", "5"],
    ]
    for r, row in enumerate(rubric.rows):
        for c, cell in enumerate(row.cells):
            set_cell(cell, rows[r][c], bold=(r == 0), size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if c != 1 else WD_ALIGN_PARAGRAPH.LEFT)
            borders(cell)
            margins(cell, top=50, bottom=50)
    add_para(doc, "#학번,이름필작성 #모든항목반드시입력 #글자크기10포인트 #글꼴 굴림 #교과세특반영될수있음", size=8.5, after=4)

    q_table = doc.add_table(rows=len(SELF_ANSWERS) + 1, cols=2)
    table_width(q_table)
    for i, (question, answer) in enumerate(SELF_ANSWERS):
        set_cell(q_table.rows[i].cells[0], question, True, 9.5)
        set_cell(q_table.rows[i].cells[1], answer, False, 9.5)
        for cell in q_table.rows[i].cells:
            borders(cell)
            margins(cell, top=70, bottom=70)
    set_cell(q_table.rows[-1].cells[0], "1번~6번 내용을\n종합하여 작성하세요.\n(12줄 내외)", True, 9.5)
    set_cell(q_table.rows[-1].cells[1], SELF_SYNTHESIS, False, 9.5)
    for cell in q_table.rows[-1].cells:
        borders(cell)
        margins(cell, top=80, bottom=80)
    doc.save(SELF_OUT)


if __name__ == "__main__":
    build_report()
    build_self_eval()
